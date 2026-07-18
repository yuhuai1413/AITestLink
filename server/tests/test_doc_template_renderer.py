from __future__ import annotations

import base64
from types import SimpleNamespace

from docx import Document

from app.services.doc_template_parser import parse_docx_template
from app.services.doc_template_renderer import render_docx_from_template


def _build_template(path):
    doc = Document()
    doc.add_paragraph("[软件名称]")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "需求标识"
    table.rows[0].cells[1].text = "测试项"
    table.rows[0].cells[2].text = "测试类型"
    table.rows[1].cells[0].text = "[需求标识]"
    table.rows[1].cells[1].text = "[测试项]"
    table.rows[1].cells[2].text = "[测试类型]"
    case_table = doc.add_table(rows=2, cols=5)
    case_table.rows[0].cells[0].text = "测试用例标识"
    case_table.rows[0].cells[1].text = "预期结果"
    case_table.rows[0].cells[2].text = "实测结果"
    case_table.rows[0].cells[3].text = "结果判定"
    case_table.rows[0].cells[4].text = "问题报告单标识/备注"
    doc.save(path)


def test_docx_template_parser_detects_tables(tmp_path):
    template_path = tmp_path / "template.docx"
    _build_template(template_path)

    structure = parse_docx_template(template_path)

    assert structure["fileName"] == "template.docx"
    assert structure["tableKinds"]["test_points"] == 1
    assert structure["tableKinds"]["test_cases"] == 1
    assert "[软件名称]" in structure["placeholders"]


def test_docx_template_renderer_fills_template_without_ai(tmp_path):
    template_path = tmp_path / "template.docx"
    _build_template(template_path)
    project = SimpleNamespace(id="p1", name="文件上传测试项目", test_type="功能测试", description="")
    requirements = [
        SimpleNamespace(id="req-uuid-1", req_id="REQ_001", module="文件管理", feature="文件上传", rule="支持上传文件", risk="高")
    ]
    points = [
        SimpleNamespace(
            id="tp-uuid-1",
            point_code="TP_001",
            requirement_id="req-uuid-1",
            module="文件管理",
            type="正常流程",
            title="文件上传成功",
            description="合法文件上传成功",
            priority="P1",
        )
    ]
    cases = [
        SimpleNamespace(
            case_code="TC_001",
            requirement_id="req-uuid-1",
            test_point_id="tp-uuid-1",
            module="文件管理",
            feature="文件上传",
            title="上传合法文件",
            expected_result="上传成功",
            actual_result="",
            passed="未执行",
            remark="",
        )
    ]

    result = render_docx_from_template(
        template_path,
        template_name="软件测试报告",
        project=project,
        requirements=requirements,
        test_points=points,
        test_cases=cases,
    )
    generated_path = tmp_path / "generated.docx"
    generated_path.write_bytes(base64.b64decode(result["docxBase64"]))
    generated = Document(generated_path)

    assert generated.paragraphs[0].text == "文件上传测试项目"
    assert [cell.text for cell in generated.tables[0].rows[1].cells] == ["REQ_001", "文件上传成功", "正常流程"]
    assert [cell.text for cell in generated.tables[1].rows[1].cells] == ["TC_001", "上传成功", "", "未执行", ""]
