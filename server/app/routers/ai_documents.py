"""AI document generation endpoints and background workflow."""

import json
import logging
import uuid

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from pydantic import BaseModel as _BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import async_session, get_db
from app.models.ai_task import AITask
from app.models.requirement import Requirement
from app.models.test_point import TestPoint
from app.models.test_case import TestCase
from app.routers.auth import get_current_user
from app.services.ai_service import AIService, check_config_for_task
from app.services.ai_input_builder import document_context, validate_case_runtime_fields, validate_persisted_traceability
from app.services.ai_task_support import friendly_error as _friendly_error, update_task_status as _update_task_status
from app.utils import model_to_dict, verify_project_owner

logger = logging.getLogger(__name__)
router = APIRouter()
ai_service = AIService()

# ─── 文档生成 ───

class GenerateDocsRequest(_BaseModel):
    template_id: str | None = None  # tpl-plan / tpl-spec / ... None=全部


async def run_generate_docs(task_id: str, project_id: str, user_id: str, template_id: str | None = None):
    async with async_session() as db:
        try:
            from app.models.project import Project
            from app.models.doc_template import DocTemplate

            # 获取项目信息
            proj_result = await db.execute(select(Project).where(Project.id == project_id))
            project = proj_result.scalar_one_or_none()

            project_info = f"项目名称:{project.name if project else '未知'}"

            # 获取需求
            req_result = await db.execute(
                select(Requirement).where(Requirement.project_id == project_id)
            )
            requirements = req_result.scalars().all()

            # 获取测试点
            tp_result = await db.execute(
                select(TestPoint).where(TestPoint.project_id == project_id)
            )
            points = tp_result.scalars().all()

            # 获取测试用例
            tc_result = await db.execute(
                select(TestCase).where(TestCase.project_id == project_id)
            )
            cases = tc_result.scalars().all()
            validate_persisted_traceability(points, cases)
            validate_case_runtime_fields(cases)
            req_text, tp_text, tc_text = document_context(requirements, points, cases)

            # 从数据库读取模板配置
            if template_id:
                tpl_result = await db.execute(
                    select(DocTemplate).where(
                        DocTemplate.config_key == template_id,
                        DocTemplate.user_id == user_id,
                    )
                )
                tpl = tpl_result.scalar_one_or_none()
                tpl_list = [tpl] if tpl else []
            else:
                tpl_result = await db.execute(
                    select(DocTemplate).where(DocTemplate.user_id == user_id)
                )
                tpl_list = tpl_result.scalars().all()

            # 逐个模板生成
            results = []
            for tpl in tpl_list:
                # 使用模板专属 prompt
                custom_prompt = tpl.prompt_template or ""
                if custom_prompt and custom_prompt != "test":
                    # 将模板 prompt 作为 system_prompt，项目数据作为 user_prompt
                    doc_result = await ai_service.generate_doc_by_template(
                        custom_prompt, project_info, req_text, tp_text, tc_text, user_id
                    )
                else:
                    # 使用默认 prompt
                    doc_result = await ai_service.generate_test_documents(
                        project_info, req_text, tp_text, tc_text, user_id
                    )

                # 尝试读取 Word 模板文件并合并内容
                if tpl.template_file:
                    template_path = os.path.join("uploads", "doc-templates", tpl.template_file)
                    if os.path.exists(template_path):
                        try:
                            doc_result = _merge_docx_template(template_path, doc_result, project.name or "")
                        except Exception as merge_err:
                            logger.warning(f"Failed to merge docx template: {merge_err}")

                results.append({
                    "templateId": tpl.config_key,
                    "templateName": tpl.name,
                    **(doc_result if isinstance(doc_result, dict) else {"content": str(doc_result)}),
                })

            # 保存结果到 AITask
            import json
            task_result = await db.execute(select(AITask).where(AITask.id == task_id))
            task = task_result.scalar_one_or_none()
            if task:
                task.result = json.dumps(results if len(results) > 1 else (results[0] if results else {}), ensure_ascii=False)

            await _update_task_status(db, task_id, "成功")
        except Exception as e:
            logger.exception("run_generate_docs failed")
            await _update_task_status(db, task_id, "失败", _friendly_error(e, "文档生成"))


def _merge_docx_template(template_path: str, doc_result: dict, project_name: str) -> dict:
    """读取 Word 模板，替换占位符，返回更新后的 result"""
    try:
        from docx import Document
        doc = Document(template_path)
        content = doc_result.get("content", "")

        # 替换段落中的占位符
        for para in doc.paragraphs:
            if "[软件名称]" in para.text:
                for run in para.runs:
                    if "[软件名称]" in run.text:
                        run.text = run.text.replace("[软件名称]", project_name)

        # 将 AI 生成的内容追加到文档末尾
        if content:
            doc.add_paragraph("")
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                try:
                    if stripped.startswith("## "):
                        doc.add_heading(stripped[3:].strip(), level=2)
                    elif stripped.startswith("# "):
                        doc.add_heading(stripped[2:].strip(), level=1)
                    elif stripped.startswith("- "):
                        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
                    else:
                        doc.add_paragraph(stripped)
                except Exception:
                    # 样式不存在时用默认样式
                    doc.add_paragraph(stripped)

        # 保存到临时路径
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()

        # 读取为 base64 返回
        import base64
        with open(tmp.name, "rb") as f:
            docx_bytes = f.read()
        os.unlink(tmp.name)

        doc_result["docxBase64"] = base64.b64encode(docx_bytes).decode()
        doc_result["docxFileName"] = f"{project_name}-{doc_result.get('title', '文档')}.docx"
        return doc_result
    except Exception as e:
        logger.warning(f"_merge_docx_template failed: {e}")
        return doc_result


@router.post("/projects/{project_id}/ai/generate-docs")
async def generate_docs(
    project_id: str,
    body: GenerateDocsRequest | None = None,
    background_tasks: BackgroundTasks = None,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    config_check = await check_config_for_task("文档生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    template_id = body.template_id if body else None

    task = AITask(
        id=str(uuid.uuid4()),
        project_id=project_id,
        type="文档生成",
        status="执行中",
        model_name="AI",
    )
    db.add(task)
    await db.commit()

    background_tasks.add_task(run_generate_docs, task.id, project_id, user["sub"], template_id)
    return model_to_dict(task)
