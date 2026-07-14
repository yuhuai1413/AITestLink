import json
import logging
import os
import uuid
from datetime import datetime, timezone

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import async_session, get_db
from app.models.ai_task import AITask
from app.models.file_asset import FileAsset
from app.models.status_log import StatusLog
from app.models.project import Project
from app.models.requirement import Requirement
from app.models.test_point import TestPoint
from app.models.test_case import TestCase
from app.routers.auth import get_current_user
from app.routers.deps import get_current_user_sse
from app.services.ai_service import AIService, check_config_for_task
from app.utils import model_to_dict
from app.utils import verify_project_owner

logger = logging.getLogger(__name__)


def _friendly_error(err: Exception, task_type: str = "") -> str:
    """将后端异常转换为中文用户友好提示"""
    import httpx as _httpx
    msg = str(err)

    # ── 按异常类型匹配（优先级最高，str() 可能为空） ──
    if isinstance(err, _httpx.TimeoutException):
        return "模型服务响应超时，请检查网络连接后重试"
    if isinstance(err, _httpx.ConnectError):
        return "无法连接到模型服务，请检查网络连接或模型配置中的 API 地址"
    if isinstance(err, _httpx.ConnectTimeout):
        return "模型服务连接超时，请检查网络连接或模型配置中的 API 地址"
    if isinstance(err, _httpx.ReadTimeout):
        return "模型服务响应超时，请稍后重试（模型生成内容较多时可能需要更长时间）"
    if isinstance(err, _httpx.WriteTimeout):
        return "模型服务发送数据超时，请稍后重试"
    if isinstance(err, _httpx.PoolTimeout):
        return "模型服务连接池超时，请稍后重试"

    # ── httpx HTTPStatusError: 检查 status_code 属性 ──
    status_code = getattr(err, "response", None)
    if status_code is not None:
        code = getattr(status_code, "status_code", 0)
        if code == 400:
            return f"{task_type or '请求'}参数错误（400），请检查输入数据后重试"
        if code == 401 or code == 403:
            return "API Key 无效或无权限，请在模型配置中检查 API Key"
        if code == 404:
            return "模型服务地址不可用（404），请在模型配置中检查 API 地址是否正确"
        if code == 429:
            return "模型服务请求过于频繁（限流），请稍后重试"
        if code >= 500:
            return "模型服务暂时不可用，请稍后重试"

    # ── 字符串匹配兜底 ──
    if "404" in msg or "Not Found" in msg:
        return "模型服务地址不可用（404），请在模型配置中检查 API 地址是否正确"
    if "401" in msg or "Unauthorized" in msg or "403" in msg or "Forbidden" in msg:
        return "API Key 无效或无权限，请在模型配置中检查 API Key"
    if "429" in msg or "Too Many Requests" in msg or "Rate limit" in msg:
        return "模型服务请求过于频繁（限流），请稍后重试"
    if "500" in msg or "502" in msg or "503" in msg or "Internal Server" in msg or "Bad Gateway" in msg or "Service Unavailable" in msg:
        return "模型服务暂时不可用，请稍后重试"
    if "Connect" in msg or "connect" in msg or "ConnectionRefused" in msg:
        return "无法连接到模型服务，请检查网络连接或模型配置中的 API 地址"
    if "Timeout" in msg or "timeout" in msg:
        return "模型服务响应超时，请检查网络连接后重试"
    if "SSL" in msg or "ssl" in msg:
        return "SSL 连接错误，请检查网络环境"
    if "JSONDecodeError" in msg or "json" in msg.lower():
        return "模型返回的数据格式异常，无法解析，请稍后重试"
    if "KeyError" in msg:
        return "模型返回数据结构异常，请稍后重试"

    # ── AI 配置相关 ──
    if "配置不存在" in msg or "模型配置" in msg:
        return msg
    if "请先" in msg:
        return msg

    # ── 默认兜底 ──
    label = task_type or "任务"
    return f"{label}失败：{msg[:200]}" if msg else f"{label}失败：未知错误，请稍后重试"



router = APIRouter()
ai_service = AIService()


# ─── 工具函数 ───

async def _update_task_status(db: AsyncSession, task_id: str, status: str, error: str | None = None):
    result = await db.execute(select(AITask).where(AITask.id == task_id))
    task = result.scalar_one_or_none()
    if task:
        task.status = status
        task.finished_at = datetime.now(timezone.utc)
        if error:
            task.error_message = error[:2000]
        await db.commit()


# 中文模块名 -> 英文缩写映射
_MODULE_ABBR: dict[str, str] = {
    # 用户与权限
    "用户管理": "USER", "用户": "USER", "登录": "LOGIN", "注销": "LOGOUT",
    "注册": "SIGNUP", "权限": "PERM", "角色": "ROLE", "账号": "ACCOUNT",
    "密码": "PWD", "忘记密码": "FORGOT", "退出": "LOGOUT", "切换": "SWITCH",
    "个人": "PERSONAL", "头像": "AVATAR", "邮箱": "EMAIL", "手机": "PHONE",
    "短信": "SMS", "验证码": "CODE", "实名": "VERIFY", "认证": "AUTH",
    # 组织架构
    "部门": "DEPT", "组织": "ORG", "公司": "COMP", "团队": "TEAM",
    "员工": "STAFF", "职位": "POSITION", "职级": "LEVEL",
    # 客户与销售
    "客户管理": "CUST", "客户": "CUST", "线索": "LEAD", "商机": "OPP",
    "合同": "CONTRACT", "报价": "QUOTE", "销售": "SALES", "回款": "PAYMENT",
    # 订单与交易
    "订单": "ORDER", "交易": "TRADE", "支付": "PAY", "退款": "REFUND",
    "购物车": "CART", "结算": "CHECKOUT",
    # 商品与库存
    "商品": "GOODS", "产品": "PROD", "库存": "STOCK", "SKU": "SKU",
    "分类": "CATEGORY", "品牌": "BRAND", "供应商": "SUPPLIER",
    "采购": "PURCHASE", "入库": "INBOUND", "出库": "OUTBOUND",
    "盘点": "INVENTORY",
    # 财务
    "财务": "FINANCE", "发票": "INVOICE", "报销": "REIMBURSE",
    "预算": "BUDGET", "账单": "BILL", "对账": "RECONCILE",
    # 审批与流程
    "审批": "APPROVE", "流程": "FLOW", "工作流": "WORKFLOW",
    "请假": "LEAVE", "考勤": "ATTEND", "打卡": "CHECKIN",
    # 通知与消息
    "通知": "NOTIF", "消息": "MSG", "公告": "ANNOUNCE", "提醒": "REMIND",
    # 报表与统计
    "报表": "REPORT", "统计": "STAT", "分析": "ANALYSIS", "图表": "CHART",
    "仪表盘": "DASH", "数据汇总": "SUMMARY", "导出": "EXPORT", "导入": "IMPORT",
    # 文档与文件
    "文件": "FILE", "文档": "DOC", "模板": "TPL", "附件": "ATTACH",
    "上传": "UPLOAD", "下载": "DL", "预览": "PREVIEW",
    # 系统与配置
    "系统": "SYS", "设置": "SETTING", "配置": "CONFIG", "参数": "PARAM",
    "字典": "DICT", "日志": "LOG", "操作日志": "OPLOG", "登录日志": "LOGINLOG",
    # 搜索与筛选
    "搜索": "SEARCH", "筛选": "FILT", "排序": "SORT",
    # UI 组件
    "表格": "TABLE", "表单": "FORM", "弹窗": "DLG", "分页": "PAGE",
    "列表": "LIST", "详情": "DETAIL", "编辑": "EDIT", "新增": "ADD",
    "删除": "DEL", "导航": "NAV", "首页": "HOME", "面包屑": "BREADCRUMB",
    # 测试相关
    "测试": "TEST", "缺陷": "BUG", "需求": "REQ", "脚本": "SCRIPT",
    "用例": "CASE", "执行": "EXEC", "评审": "REVIEW", "生成": "GEN",
    "回归": "REG", "冒烟": "SMOKE", "接口": "API", "性能": "PERF",
    "安全": "SEC", "兼容": "COMPAT", "自动化": "AUTO",
    # 通用
    "核心": "CORE", "基础": "BASE", "公共": "COMMON", "通用": "COMMON",
    "状态": "STATUS", "类型": "TYPE", "标签": "TAG", "分类管理": "CATMGR",
    "版本": "VERSION", "发布": "RELEASE", "部署": "DEPLOY",
    "缓存": "CACHE", "任务": "TASK", "调度": "SCHEDULE",
    "字权限": "DATAPERM", "数据权限": "DATAPERM",
}

# 已使用的英文缩写，防止重复
_used_abbrs: dict[str, int] = {}


def _to_eng_abbr(module: str) -> str:
    """将中文模块名转为英文缩写，确保输出纯英文+数字，不含中文。"""
    if module in _MODULE_ABBR:
        return _MODULE_ABBR[module]
    # 尝试拆分组合词：如 "合同审批" → "合同"+"审批" → "CONTRACT"+"APPROVE"
    for cn, en in _MODULE_ABBR.items():
        if len(cn) >= 2 and cn in module:
            # 找到子串匹配，用对应的英文缩写拼接
            remaining = module.replace(cn, "", 1)
            rest_abbr = _MODULE_ABBR.get(remaining, "") if remaining else ""
            if rest_abbr:
                return f"{en}_{rest_abbr}"
            return en
    # 无法匹配时，取每个汉字拼音首字母（用 Unicode 区间粗略判断）
    # 退而求其次：取模块名中每个非中文字符的大写
    ascii_chars = [c.upper() for c in module if c.isascii() and c.isalpha()]
    if ascii_chars:
        return "".join(ascii_chars[:4])
    # 全是中文且无匹配，用 MOD + 序号避免重复
    return "MOD"


def _module_counter(items: list[dict], prefix: str) -> list[tuple[str, str]]:
    """为按模块分组的项目生成编号。返回 [(id, module), ...]
    编号格式: PREFIX_ABBR_NNN（纯英文+数字，不含中文）
    """
    module_counters: dict[str, int] = {}
    result = []
    for item in items:
        module = item["module"]
        abbr = _to_eng_abbr(module)
        module_counters[abbr] = module_counters.get(abbr, 0) + 1
        result.append((f"{prefix}_{abbr}_{module_counters[abbr]:03d}", module))
    return result


def _read_file_content(file_obj: FileAsset) -> str:
    """读取单个文件的内容，返回文本。"""
    if not file_obj.storage_path or not os.path.exists(file_obj.storage_path):
        return ""

    ext = file_obj.name.rsplit(".", 1)[-1].lower() if "." in file_obj.name else ""

    try:
        if ext in ("docx", "doc"):
            from docx import Document
            doc = Document(file_obj.storage_path)
            parts = []
            # 读取段落
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            # 读取表格内容
            for i, table in enumerate(doc.tables):
                table_text = [f"[表格{i+1}]"]
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_text.append(" | ".join(cells))
                if len(table_text) > 1:
                    parts.append("\n".join(table_text))
            text = "\n".join(parts)
            return text

        if ext in ("xlsx", "xls"):
            import openpyxl
            wb = openpyxl.load_workbook(file_obj.storage_path, read_only=True, data_only=True)
            parts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    parts.append(f"Sheet: {sheet}\n" + "\n".join(rows))
            wb.close()
            return "\n".join(parts)

        if ext == "pdf":
            import PyPDF2
            reader = PyPDF2.PdfReader(file_obj.storage_path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text

        if ext in ("md", "txt", "json", "yaml", "yml", "csv"):
            with open(file_obj.storage_path, "r", errors="ignore") as fh:
                return fh.read()

        return f"(不支持的格式: {ext})"

    except Exception as e:
        return f"(读取失败: {str(e)[:100]})"


# ─── 后台任务 ───

async def run_parse_requirements(task_id: str, project_id: str, file_content: str, user_id: str):
    async with async_session() as db:
        try:
            file_result = await db.execute(
                select(FileAsset).where(FileAsset.project_id == project_id)
            )
            files = file_result.scalars().all()

            for f in files:
                f.parse_status = "解析中"
            await db.commit()

            # 删除旧需求，避免重复追加
            from sqlalchemy import delete
            await db.execute(delete(Requirement).where(Requirement.project_id == project_id))
            await db.commit()

            logger.info(f"run_parse_requirements: file_content_len={len(file_content)}, user_id={user_id}")
            requirements = await ai_service.parse_requirements(file_content, user_id)
            logger.info(f"parse_requirements result: type={type(requirements).__name__}, len={len(requirements) if isinstance(requirements, (list, dict)) else 'N/A'}, preview={str(requirements)[:200]}")

            # 确保返回的是列表；如果 LLM 返回了 dict，尝试从各种常见结构中提取列表
            if isinstance(requirements, dict):
                for v in requirements.values():
                    if isinstance(v, list) and v and isinstance(v[0], dict) and "module" in v[0]:
                        requirements = v
                        break
                if isinstance(requirements, dict):
                    for k, v in requirements.items():
                        if isinstance(v, list) and v and isinstance(v[0], dict):
                            requirements = v
                            logger.info(f"extracted list from dict key '{k}', len={len(v)}")
                            break
            # 如果返回单个 dict（截断响应），尝试包装为列表
            if isinstance(requirements, dict):
                if {"module", "feature"}.issubset(requirements.keys()):
                    requirements = [requirements]
                else:
                    logger.error(f"parse_requirements: invalid result, type={type(requirements).__name__}, preview={str(requirements)[:500]}")
                    for f in files:
                        f.parse_status = "失败"
                        f.parse_error = "AI 未返回有效的解析数据"
                    await db.commit()
                    await _update_task_status(db, task_id, "失败", "AI 未返回有效的解析数据，请检查模型配置或文件内容后重试")
                    return
            if not isinstance(requirements, list) or not requirements:
                logger.error(f"parse_requirements: invalid result, type={type(requirements).__name__}, preview={str(requirements)[:500]}")
                for f in files:
                    f.parse_status = "失败"
                    f.parse_error = "AI 未返回有效的解析数据"
                await db.commit()
                await _update_task_status(db, task_id, "失败", "AI 未返回有效的解析数据，请检查模型配置或文件内容后重试")
                return

            for req in requirements:
                db.add(Requirement(
                    id=str(uuid.uuid4()),
                    project_id=project_id,
                    module=req["module"],
                    feature=req["feature"],
                    source=req.get("source", ""),
                    risk=req.get("risk", "中"),
                    rule=req.get("rule", ""),
                    question=req.get("question", ""),
                ))

            for f in files:
                f.parse_status = "已完成"
                f.parse_error = ""

            await db.commit()
            await _update_task_status(db, task_id, "成功")

        except Exception as e:
            error_msg = _friendly_error(e, "需求解析")
            logger.exception("run_parse_requirements failed")
            try:
                file_result = await db.execute(
                    select(FileAsset).where(FileAsset.project_id == project_id)
                )
                for f in file_result.scalars().all():
                    f.parse_status = "失败"
                    f.parse_error = error_msg
                
                # 记录解析失败日志（不改变test_status，保持"测试中"状态）
                
                await db.commit()
            except Exception:
                logger.exception("Failed to update file status on error")
            await _update_task_status(db, task_id, "失败", _friendly_error(e, "需求解析"))


async def run_generate_test_points(task_id: str, project_id: str, user_id: str):
    logger.info(f"run_generate_test_points: task_id={task_id}, project_id={project_id}, user_id={user_id}")
    async with async_session() as db:
        try:
            result = await db.execute(
                select(Requirement).where(Requirement.project_id == project_id)
            )
            requirements = result.scalars().all()
            if not requirements:
                await _update_task_status(db, task_id, "失败", "需求列表为空，无法生成测试点")
                return
            req_text = "\n".join(
                f"- 模块:{r.module} 功能:{r.feature} 规则:{r.rule}" for r in requirements
            )
            # 删除旧测试点，避免重复追加
            from sqlalchemy import delete
            await db.execute(delete(TestPoint).where(TestPoint.project_id == project_id))
            await db.commit()

            points = await ai_service.generate_test_points(req_text, user_id)
            logger.info(f"generate_test_points: type={type(points).__name__}, len={len(points) if isinstance(points, (list, dict)) else 'N/A'}")

            # 确保返回的是列表；如果 LLM 返回了 dict，尝试从各种常见结构中提取列表
            if isinstance(points, dict):
                for v in points.values():
                    if isinstance(v, list) and v and isinstance(v[0], dict) and "module" in v[0]:
                        points = v
                        break
                if isinstance(points, dict):
                    for k, v in points.items():
                        if isinstance(v, list) and v and isinstance(v[0], dict):
                            points = v
                            logger.info(f"extracted list from dict key '{k}', len={len(v)}")
                            break
            # 如果返回单个 dict（截断响应），尝试包装为列表
            if isinstance(points, dict):
                if {"module", "type", "title"}.issubset(points.keys()):
                    points = [points]
                else:
                    logger.error(f"generate_test_points: invalid result, type={type(points).__name__}, preview={str(points)[:500]}")
                    await _update_task_status(db, task_id, "失败", "AI 未返回有效的测试点数据，请检查模型配置或需求数据后重试")
                    return
            if not isinstance(points, list) or not points:
                logger.error(f"generate_test_points: invalid result, type={type(points).__name__}, preview={str(points)[:500]}")
                await _update_task_status(db, task_id, "失败", "AI 未返回有效的测试点数据，请检查模型配置或需求数据后重试")
                return

            for pt in points:
                db.add(TestPoint(
                    id=pt.get("id") or str(uuid.uuid4()),
                    project_id=project_id,
                    module=pt["module"],
                    type=pt["type"],
                    title=pt["title"],
                    description=pt.get("description", ""),
                    priority=pt.get("priority", "P1"),
                    automatable=bool(pt.get("automatable", False)),
                ))


            await db.commit()
            logger.info(f"run_generate_test_points: task {task_id} completed successfully, {len(points)} points saved")
            await _update_task_status(db, task_id, "成功")

        except Exception as e:
            logger.exception(f"run_generate_test_points failed: task_id={task_id}")
            await _update_task_status(db, task_id, "失败", _friendly_error(e, "测试点生成"))


async def run_generate_test_cases(task_id: str, project_id: str, user_id: str):
    async with async_session() as db:
        try:
            tp_result = await db.execute(
                select(TestPoint).where(TestPoint.project_id == project_id)
            )
            points = tp_result.scalars().all()
            pt_text = "\n".join(
                f"- 模块:{tp.module} 类型:{tp.type} 标题:{tp.title} 优先级:{tp.priority}"
                for tp in points
            )
            # 删除旧测试用例，避免重复追加
            from sqlalchemy import delete
            await db.execute(delete(TestCase).where(TestCase.project_id == project_id))
            await db.commit()

            cases = await ai_service.generate_test_cases(pt_text, user_id)
            logger.info(f"generate_test_cases: type={type(cases).__name__}, len={len(cases) if isinstance(cases, (list, dict)) else 'N/A'}")

            # 确保返回的是列表；如果 LLM 返回了 dict，尝试从各种常见结构中提取列表
            if isinstance(cases, dict):
                # 方式1: 直接检查第一个值是否为 dict list
                for v in cases.values():
                    if isinstance(v, list) and v and isinstance(v[0], dict) and "module" in v[0]:
                        cases = v
                        break
                # 方式2: key 名包含 cases/test_cases/testcase 的 dict
                if isinstance(cases, dict):
                    for k, v in cases.items():
                        if isinstance(v, list) and v and isinstance(v[0], dict):
                            cases = v
                            logger.info(f"extracted list from dict key '{k}', len={len(v)}")
                            break
            # 如果返回单个 dict（截断响应），尝试包装为列表
            if isinstance(cases, dict):
                if {"module", "title"}.issubset(cases.keys()):
                    cases = [cases]
                else:
                    logger.error(f"generate_test_cases: invalid result, type={type(cases).__name__}, preview={str(cases)[:500]}")
                    await _update_task_status(db, task_id, "失败", "AI 未返回有效的测试用例数据，请检查模型配置或测试点数据后重试")
                    return
            if not isinstance(cases, list) or not cases:
                logger.error(f"generate_test_cases: invalid result, type={type(cases).__name__}, preview={str(cases)[:500]}")
                await _update_task_status(db, task_id, "失败", "AI 未返回有效的测试用例数据，请检查模型配置或测试点数据后重试")
                return

            for (case_code, _), c in zip(_module_counter(cases, "TC"), cases):
                db.add(TestCase(
                    id=str(uuid.uuid4()),
                    project_id=project_id,
                    case_code=case_code,
                    module=c["module"],
                    feature=c.get("feature", ""),
                    title=c["title"],
                    priority=c.get("priority", "P1"),
                    precondition=c.get("precondition", ""),
                    steps=c.get("steps", ""),
                    test_data=json.dumps(c.get("testData", ""), ensure_ascii=False) if isinstance(c.get("testData"), (dict, list)) else str(c.get("testData", "")),
                    expected_result=c.get("expectedResult", ""),
                    test_type=c.get("testType", "功能测试"),
                    automation=c.get("automation", "待评估"),
                ))


            await db.commit()
            await _update_task_status(db, task_id, "成功")

        except Exception as e:
            logger.exception("run_generate_test_cases failed")
            friendly = _friendly_error(e, "用例生成")
            await _update_task_status(db, task_id, "失败", friendly)


# ─── 路由 ───

@router.get("/projects/{project_id}/ai/tasks")
async def list_ai_tasks(
    project_id: str,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])
    result = await db.execute(
        select(AITask).where(AITask.project_id == project_id).order_by(AITask.created_at.desc())
    )
    return [model_to_dict(t) for t in result.scalars().all()]


@router.get("/projects/{project_id}/ai/check-config/{task_type}")
async def check_ai_config(
    project_id: str,
    task_type: str,
    user: dict = Depends(get_current_user),
):
    """检查用户是否已配置指定 AI 任务的模型"""
    result = await check_config_for_task(task_type, user["sub"])
    return result


@router.post("/projects/{project_id}/ai/parse-requirements")
async def parse_requirements(
    project_id: str,
    background_tasks: BackgroundTasks,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])
    # 检查配置
    config_check = await check_config_for_task("需求解析", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    file_result = await db.execute(
        select(FileAsset).where(FileAsset.project_id == project_id)
    )
    files = file_result.scalars().all()
    if not files:
        return {"error": "No files found. Please upload files first."}

    content_parts = []
    # 按文件类型排序：需求文档(docx)优先，辅助文档(xlsx/txt等)其次
    def _file_priority(f):
        ext = f.name.rsplit(".", 1)[-1].lower() if "." in f.name else ""
        if ext in ("docx", "doc"):
            return 0  # 需求文档优先
        return 1  # 辅助文档其次

    sorted_files = sorted(files, key=_file_priority)
    for f in sorted_files:
        text = _read_file_content(f)
        if text:
            content_parts.append(f"[{f.name}]\n{text}")

    file_content = "\n---\n".join(content_parts) if content_parts else "No readable content"

    task = AITask(
        id=str(uuid.uuid4()),
        project_id=project_id,
        type="需求解析",
        status="执行中",
        model_name="AI",
    )
    db.add(task)
    await db.commit()

    background_tasks.add_task(run_parse_requirements, task.id, project_id, file_content, user["sub"])
    return model_to_dict(task)


@router.post("/projects/{project_id}/ai/generate-test-points")
async def generate_test_points(
    project_id: str,
    background_tasks: BackgroundTasks,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])
    # 检查是否有需求数据
    req_result = await db.execute(
        select(Requirement).where(Requirement.project_id == project_id)
    )
    if not req_result.scalars().first():
        raise HTTPException(status_code=400, detail="需求列表为空，请先在「需求列表」页面完成需求解析")

    # 检查配置
    config_check = await check_config_for_task("测试点生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    task = AITask(
        id=str(uuid.uuid4()),
        project_id=project_id,
        type="测试点生成",
        status="执行中",
        model_name="AI",
    )
    db.add(task)
    await db.commit()

    background_tasks.add_task(run_generate_test_points, task.id, project_id, user["sub"])
    return model_to_dict(task)


@router.post("/projects/{project_id}/ai/generate-test-cases")
async def generate_test_cases(
    project_id: str,
    background_tasks: BackgroundTasks,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])
    # 检查测试点是否存在
    tp_result = await db.execute(
        select(TestPoint).where(TestPoint.project_id == project_id)
    )
    if not tp_result.scalars().first():
        raise HTTPException(status_code=400, detail="测试点列表为空，请先在「测试点」页面生成测试点")
    # 检查配置
    config_check = await check_config_for_task("用例生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    task = AITask(
        id=str(uuid.uuid4()),
        project_id=project_id,
        type="用例生成",
        status="执行中",
        model_name="AI",
    )
    db.add(task)
    await db.commit()

    background_tasks.add_task(run_generate_test_cases, task.id, project_id, user["sub"])
    return model_to_dict(task)


# ─── 执行脚本分析 ───

async def run_execute_scripts_analysis(task_id: str, project_id: str, user_id: str):
    async with async_session() as db:
        try:
            from app.models.automation_script import AutomationScript
            result = await db.execute(
                select(AutomationScript).where(AutomationScript.project_id == project_id)
            )
            scripts = result.scalars().all()
            if not scripts:
                await _update_task_status(db, task_id, "失败", "自动化脚本列表为空，无法分析")
                return

            scripts_text = "\n".join(
                f"- ID:{s.id} 类型:{s.script_type} 框架:{s.framework} 语言:{s.language} "
                f"状态:{s.status} 代码片段:{(s.code or '')[:200]}"
                for s in scripts
            )

            execution_results = "当前脚本状态统计：" + "\n".join(
                f"- {s.status}: {sum(1 for x in scripts if x.status == s.status)} 个"
                for s in scripts
            )

            analysis = await ai_service.analyze_script_execution(scripts_text, execution_results, user_id)

            task_result = await db.execute(select(AITask).where(AITask.id == task_id))
            task = task_result.scalar_one_or_none()
            if task:
                import json
                task.result = json.dumps(analysis, ensure_ascii=False) if isinstance(analysis, dict) else str(analysis)

            await _update_task_status(db, task_id, "成功")
        except Exception as e:
            logger.exception("run_execute_scripts_analysis failed")
            await _update_task_status(db, task_id, "失败", _friendly_error(e, "执行脚本"))


@router.post("/projects/{project_id}/ai/execute-scripts")
async def execute_scripts(
    project_id: str,
    background_tasks: BackgroundTasks,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    from app.models.automation_script import AutomationScript
    script_result = await db.execute(
        select(AutomationScript).where(AutomationScript.project_id == project_id)
    )
    if not script_result.scalars().first():
        raise HTTPException(status_code=400, detail="自动化脚本列表为空，请先生成脚本")

    config_check = await check_config_for_task("执行脚本", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    task = AITask(
        id=str(uuid.uuid4()),
        project_id=project_id,
        type="执行脚本",
        status="执行中",
        model_name="AI",
    )
    db.add(task)
    await db.commit()

    background_tasks.add_task(run_execute_scripts_analysis, task.id, project_id, user["sub"])
    return model_to_dict(task)


# ─── 文档生成 ───

from pydantic import BaseModel as _BaseModel

class GenerateDocsRequest(_BaseModel):
    template_id: str | None = None  # tpl-plan / tpl-spec / ... None=全部


async def run_generate_docs(task_id: str, project_id: str, user_id: str, template_id: str | None = None):
    async with async_session() as db:
        try:
            from app.models.project import Project
            from app.models.doc_template import DocTemplate

            # 获取项目信息
            proj_result = await db.execute(select(Project).where(Project.id == project_id))
            project = proj_result.scalar_one_or_none()

            project_info = f"项目名称:{project.name if project else '未知'}"

            # 获取需求
            req_result = await db.execute(
                select(Requirement).where(Requirement.project_id == project_id)
            )
            requirements = req_result.scalars().all()
            req_text = "\n".join(
                f"- 模块:{r.module} 功能:{r.feature} 规则:{r.rule}" for r in requirements
            ) if requirements else "暂无需求"

            # 获取测试点
            tp_result = await db.execute(
                select(TestPoint).where(TestPoint.project_id == project_id)
            )
            points = tp_result.scalars().all()
            tp_text = "\n".join(
                f"- 编号:{tp.id} 模块:{tp.module} 标题:{tp.title} 优先级:{tp.priority}"
                for tp in points
            ) if points else "暂无测试点"

            # 获取测试用例
            tc_result = await db.execute(
                select(TestCase).where(TestCase.project_id == project_id)
            )
            cases = tc_result.scalars().all()
            tc_text = "\n".join(
                f"- 编号:{c.case_code} 模块:{c.module} 标题:{c.title} 优先级:{c.priority} 自动化:{c.automation}"
                for c in cases
            ) if cases else "暂无测试用例"

            # 从数据库读取模板配置
            if template_id:
                tpl_result = await db.execute(
                    select(DocTemplate).where(
                        DocTemplate.config_key == template_id,
                        DocTemplate.user_id == user_id,
                    )
                )
                tpl = tpl_result.scalar_one_or_none()
                tpl_list = [tpl] if tpl else []
            else:
                tpl_result = await db.execute(
                    select(DocTemplate).where(DocTemplate.user_id == user_id)
                )
                tpl_list = tpl_result.scalars().all()

            # 逐个模板生成
            results = []
            for tpl in tpl_list:
                # 使用模板专属 prompt
                custom_prompt = tpl.prompt_template or ""
                if custom_prompt and custom_prompt != "test":
                    # 将模板 prompt 作为 system_prompt，项目数据作为 user_prompt
                    doc_result = await ai_service.generate_doc_by_template(
                        custom_prompt, project_info, req_text, tp_text, tc_text, user_id
                    )
                else:
                    # 使用默认 prompt
                    doc_result = await ai_service.generate_test_documents(
                        project_info, req_text, tp_text, tc_text, user_id
                    )

                # 尝试读取 Word 模板文件并合并内容
                if tpl.template_file:
                    template_path = os.path.join("uploads", "doc-templates", tpl.template_file)
                    if os.path.exists(template_path):
                        try:
                            doc_result = _merge_docx_template(template_path, doc_result, project.name or "")
                        except Exception as merge_err:
                            logger.warning(f"Failed to merge docx template: {merge_err}")

                results.append({
                    "templateId": tpl.config_key,
                    "templateName": tpl.name,
                    **(doc_result if isinstance(doc_result, dict) else {"content": str(doc_result)}),
                })

            # 保存结果到 AITask
            import json
            task_result = await db.execute(select(AITask).where(AITask.id == task_id))
            task = task_result.scalar_one_or_none()
            if task:
                task.result = json.dumps(results if len(results) > 1 else (results[0] if results else {}), ensure_ascii=False)

            await _update_task_status(db, task_id, "成功")
        except Exception as e:
            logger.exception("run_generate_docs failed")
            await _update_task_status(db, task_id, "失败", _friendly_error(e, "文档生成"))


def _merge_docx_template(template_path: str, doc_result: dict, project_name: str) -> dict:
    """读取 Word 模板，替换占位符，返回更新后的 result"""
    try:
        from docx import Document
        doc = Document(template_path)
        content = doc_result.get("content", "")

        # 替换段落中的占位符
        for para in doc.paragraphs:
            if "[软件名称]" in para.text:
                for run in para.runs:
                    if "[软件名称]" in run.text:
                        run.text = run.text.replace("[软件名称]", project_name)

        # 将 AI 生成的内容追加到文档末尾
        if content:
            doc.add_paragraph("")
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                try:
                    if stripped.startswith("## "):
                        doc.add_heading(stripped[3:].strip(), level=2)
                    elif stripped.startswith("# "):
                        doc.add_heading(stripped[2:].strip(), level=1)
                    elif stripped.startswith("- "):
                        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
                    else:
                        doc.add_paragraph(stripped)
                except Exception:
                    # 样式不存在时用默认样式
                    doc.add_paragraph(stripped)

        # 保存到临时路径
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()

        # 读取为 base64 返回
        import base64
        with open(tmp.name, "rb") as f:
            docx_bytes = f.read()
        os.unlink(tmp.name)

        doc_result["docxBase64"] = base64.b64encode(docx_bytes).decode()
        doc_result["docxFileName"] = f"{project_name}-{doc_result.get('title', '文档')}.docx"
        return doc_result
    except Exception as e:
        logger.warning(f"_merge_docx_template failed: {e}")
        return doc_result


@router.post("/projects/{project_id}/ai/generate-docs")
async def generate_docs(
    project_id: str,
    body: GenerateDocsRequest | None = None,
    background_tasks: BackgroundTasks = None,
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    config_check = await check_config_for_task("文档生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    template_id = body.template_id if body else None

    task = AITask(
        id=str(uuid.uuid4()),
        project_id=project_id,
        type="文档生成",
        status="执行中",
        model_name="AI",
    )
    db.add(task)
    await db.commit()

    background_tasks.add_task(run_generate_docs, task.id, project_id, user["sub"], template_id)
    return model_to_dict(task)


# ─── 流式生成（SSE） ───────────────────────────────────────────────

from fastapi.responses import StreamingResponse


async def _stream_generate_test_points(task_id: str, project_id: str, user_id: str):
    """流式生成测试点，每批写库后通过 SSE 推送。"""
    ai_service = AIService()
    async with async_session() as db:
        try:
            result = await db.execute(
                select(Requirement).where(Requirement.project_id == project_id)
            )
            requirements = result.scalars().all()
            if not requirements:
                yield f"data: {json.dumps({'event': 'error', 'message': '需求列表为空'})}\n\n"
                return

            req_text = "\n".join(
                f"- 模块:{r.module} 功能:{r.feature} 规则:{r.rule} 风险:{r.risk}"
                for r in requirements
            )
            user_prompt = f"请以资深测试架构师的视角，根据以下需求生成全面的测试点。\n\n需求列表：\n\n{req_text[:3000]}"

            total_saved = 0
            async for batch in ai_service.generate_stream(user_prompt, "测试点生成", user_id, batch_size=5):
                if not batch:
                    yield f"data: {json.dumps({'event': 'receiving'}, ensure_ascii=False)}\n\n"
                    continue
                # 先删除旧测试点（首次写入时）
                if total_saved == 0:
                    from sqlalchemy import delete
                    await db.execute(delete(TestPoint).where(TestPoint.project_id == project_id))
                    await db.commit()

                for pt in batch:
                    db.add(TestPoint(
                        id=pt.get("id") or str(uuid.uuid4()),
                        project_id=project_id,
                        module=pt.get("module", ""),
                        type=pt.get("type", ""),
                        title=pt.get("title", ""),
                        description=pt.get("description", ""),
                        priority=pt.get("priority", "P1"),
                        automatable=bool(pt.get("automatable", False)),
                    ))
                await db.commit()
                total_saved += len(batch)
                yield f"data: {json.dumps({'event': 'progress', 'count': total_saved, 'batch': len(batch)}, ensure_ascii=False)}\n\n"

            yield f"data: {json.dumps({'event': 'done', 'total': total_saved}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.exception(f"stream_generate_test_points failed: task_id={task_id}")
            yield f"data: {json.dumps({'event': 'error', 'message': _friendly_error(e, '测试点生成')}, ensure_ascii=False)}\n\n"


@router.get("/projects/{project_id}/ai/stream-test-points")
async def stream_test_points(
    project_id: str,
    user: dict = Depends(get_current_user_sse),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    # 检查需求数据
    req_result = await db.execute(
        select(Requirement).where(Requirement.project_id == project_id)
    )
    if not req_result.scalars().first():
        raise HTTPException(status_code=400, detail="需求列表为空，请先在「需求列表」页面完成需求解析")

    # 检查配置
    config_check = await check_config_for_task("测试点生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    return StreamingResponse(
        _stream_generate_test_points("", project_id, user["sub"]),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


async def _stream_generate_test_cases(task_id: str, project_id: str, user_id: str):
    """流式生成测试用例，每批写库后通过 SSE 推送。"""
    ai_service = AIService()
    async with async_session() as db:
        try:
            tp_result = await db.execute(
                select(TestPoint).where(TestPoint.project_id == project_id)
            )
            points = tp_result.scalars().all()
            if not points:
                yield f"data: {json.dumps({'event': 'error', 'message': '测试点列表为空，请先生成测试点'})}\n\n"
                return

            pt_text = "\n".join(
                f"- 模块:{tp.module} 类型:{tp.type} 标题:{tp.title} 优先级:{tp.priority}"
                for tp in points
            )
            user_prompt = (
                "请以资深测试工程师的视角，根据以下测试点生成可执行的测试用例。要求："
                "1. 步骤使用 步骤N: 格式，每步包含验证点 "
                "2. 预期结果包含步骤编号和 应 字 "
                "3. 数据具体化，不使用模糊描述 "
                "4. 编号格式 TC_XXX_NNN "
                "5. 用例标题只写验证目标本身"
                f"\n\n测试点列表：\n\n{pt_text[:3000]}"
            )

            total_saved = 0
            async for batch in ai_service.generate_stream(user_prompt, "用例生成", user_id, batch_size=5, max_tokens=16000):
                if not batch:
                    yield f"data: {json.dumps({'event': 'receiving'}, ensure_ascii=False)}\n\n"
                    continue
                # 先删除旧测试用例（首次写入时）
                if total_saved == 0:
                    from sqlalchemy import delete
                    await db.execute(delete(TestCase).where(TestCase.project_id == project_id))
                    await db.commit()

                module_counter: dict[str, int] = {}
                for c in batch:
                    mod = c.get("module", "DEFAULT")
                    module_counter[mod] = module_counter.get(mod, 0) + 1
                    case_code = f"TC_{mod.upper().replace(' ', '_')}_{module_counter[mod]:03d}"
                    db.add(TestCase(
                        id=str(uuid.uuid4()),
                        project_id=project_id,
                        case_code=case_code,
                        module=c.get("module", ""),
                        feature=c.get("feature", ""),
                        title=c.get("title", ""),
                        priority=c.get("priority", "P1"),
                        precondition=c.get("precondition", ""),
                        steps=c.get("steps", ""),
                        test_data=json.dumps(c.get("testData", ""), ensure_ascii=False) if isinstance(c.get("testData"), (dict, list)) else str(c.get("testData", "")),
                        expected_result=c.get("expectedResult", ""),
                        test_type=c.get("testType", "功能测试"),
                        automation=c.get("automation", "待评估"),
                    ))
                await db.commit()
                total_saved += len(batch)
                yield f"data: {json.dumps({'event': 'progress', 'count': total_saved, 'batch': len(batch)}, ensure_ascii=False)}\n\n"

            yield f"data: {json.dumps({'event': 'done', 'total': total_saved}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.exception(f"stream_generate_test_cases failed: task_id={task_id}")
            yield f"data: {json.dumps({'event': 'error', 'message': _friendly_error(e, '用例生成')}, ensure_ascii=False)}\n\n"


@router.get("/projects/{project_id}/ai/stream-test-cases")
async def stream_test_cases(
    project_id: str,
    user: dict = Depends(get_current_user_sse),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    # 检查测试点数据
    tp_result = await db.execute(
        select(TestPoint).where(TestPoint.project_id == project_id)
    )
    if not tp_result.scalars().first():
        raise HTTPException(status_code=400, detail="测试点列表为空，请先在「测试点」页面生成测试点")

    # 检查配置
    config_check = await check_config_for_task("用例生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    return StreamingResponse(
        _stream_generate_test_cases("", project_id, user["sub"]),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ─── 需求解析流式 ──────────────────────────────────────────────────

async def _stream_parse_requirements(project_id: str, user_id: str):
    """流式需求解析，每批写库后通过 SSE 推送。"""
    logger.info(f"stream_parse_requirements: project_id={project_id}, user_id={user_id}")
    ai_service = AIService()
    async with async_session() as db:
        try:
            # 读取文件内容
            file_result = await db.execute(
                select(FileAsset).where(FileAsset.project_id == project_id)
            )
            files = file_result.scalars().all()
            if not files:
                logger.warning("stream_parse_requirements: no files found")
                yield f"data: {json.dumps({'event': 'error', 'message': '没有上传文件'})}\n\n"
                return

            logger.info(f"stream_parse_requirements: found {len(files)} files")

            # 标记文件为解析中
            for f in files:
                f.parse_status = "解析中"
            await db.commit()

            # 读取所有文件内容
            content_parts = []
            for f in files:
                try:
                    from app.services.document_service import DocumentService
                    doc_service = DocumentService(db)
                    text = await doc_service.get_content(f.id)
                    if text:
                        content_parts.append(f"=== 文件: {f.name} ===\n{text}")
                except Exception as e:
                    logger.warning(f"Failed to read file {f.name}: {e}")

            if not content_parts:
                for f in files:
                    f.parse_status = "失败"
                    f.parse_error = "无法读取文件内容"
                await db.commit()
                yield f"data: {json.dumps({'event': 'error', 'message': '无法读取文件内容'})}\n\n"
                return

            file_content = "\n\n".join(content_parts)
            logger.info(f"stream_parse_requirements: file_content length={len(file_content)}")

            # 删除旧需求
            from sqlalchemy import delete
            await db.execute(delete(Requirement).where(Requirement.project_id == project_id))
            await db.commit()

            user_prompt = f"请对以下文档内容进行专业的需求分析。\n\n文档内容：\n\n{file_content}"
            logger.info(f"stream_parse_requirements: starting LLM stream...")

            total_saved = 0
            async for batch in ai_service.generate_stream(user_prompt, "需求解析", user_id, batch_size=5):
                if not batch:
                    # 空 batch = LLM 仍在生成中
                    yield f"data: {json.dumps({'event': 'receiving'}, ensure_ascii=False)}\n\n"
                    continue
                logger.info(f"stream_parse_requirements: received batch of {len(batch)} items")
                for req in batch:
                    db.add(Requirement(
                        id=str(uuid.uuid4()),
                        project_id=project_id,
                        module=req.get("module", ""),
                        feature=req.get("feature", ""),
                        source=req.get("source", ""),
                        risk=req.get("risk", "中"),
                        rule=req.get("rule", ""),
                        question=req.get("question", ""),
                    ))
                await db.commit()
                total_saved += len(batch)
                yield f"data: {json.dumps({'event': 'progress', 'count': total_saved, 'batch': len(batch)}, ensure_ascii=False)}\n\n"

            # 标记文件为已完成
            for f in files:
                f.parse_status = "已完成"
                f.parse_error = ""
            await db.commit()

            yield f"data: {json.dumps({'event': 'done', 'total': total_saved}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.exception(f"stream_parse_requirements failed: project_id={project_id}")
            yield f"data: {json.dumps({'event': 'error', 'message': _friendly_error(e, '需求解析')}, ensure_ascii=False)}\n\n"


@router.get("/projects/{project_id}/ai/stream-parse-requirements")
async def stream_parse_requirements(
    project_id: str,
    user: dict = Depends(get_current_user_sse),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    # 检查文件
    file_result = await db.execute(
        select(FileAsset).where(FileAsset.project_id == project_id)
    )
    if not file_result.scalars().first():
        raise HTTPException(status_code=400, detail="没有上传文件，请先在「输入资料」页面上传文件")

    # 检查配置
    config_check = await check_config_for_task("需求解析", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    return StreamingResponse(
        _stream_parse_requirements(project_id, user["sub"]),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ─── 脚本生成流式 ──────────────────────────────────────────────────

async def _stream_generate_scripts(project_id: str, user_id: str):
    """流式生成自动化脚本，每批写库后通过 SSE 推送。"""
    ai_service = AIService()
    async with async_session() as db:
        try:
            # 获取适合自动化的测试用例
            result = await db.execute(
                select(TestCase).where(
                    TestCase.project_id == project_id,
                    TestCase.automation == "适合"
                )
            )
            test_cases = result.scalars().all()
            if not test_cases:
                yield f"data: {json.dumps({'event': 'error', 'message': '没有适合自动化的测试用例'})}\n\n"
                return

            # 删除旧脚本
            from sqlalchemy import delete
            await db.execute(delete(AutomationScript).where(AutomationScript.project_id == project_id))
            await db.commit()

            # 构建测试用例文本
            tc_text = "\n".join(
                f"- 编号:{tc.case_code} 模块:{tc.module} 标题:{tc.title} 优先级:{tc.priority} "
                f"前置条件:{tc.precondition or '无'} 步骤:{tc.steps or '无'} "
                f"预期结果:{tc.expected_result or '无'}"
                for tc in test_cases
            )
            user_prompt = f"请根据以下测试用例生成可直接运行的 Playwright 自动化测试脚本。\n\n测试用例列表：\n\n{tc_text[:4000]}"

            total_saved = 0
            async for batch in ai_service.generate_stream(user_prompt, "脚本生成", user_id, batch_size=3):
                if not batch:
                    yield f"data: {json.dumps({'event': 'receiving'}, ensure_ascii=False)}\n\n"
                    continue
                for ai_script in batch:
                    tc_id = ai_script.get("testCaseId", "")
                    matched_tc = None
                    for tc in test_cases:
                        if tc.case_code == tc_id or tc.id == tc_id:
                            matched_tc = tc
                            break

                    module = (matched_tc or test_cases[0]).module or ""
                    abbr_map = {
                        "用户管理": "USER", "订单处理": "ORDER", "菜单": "MENU",
                        "客户管理": "CUST", "登录": "LOGIN", "系统": "SYS",
                        "权限": "PERM", "配置": "CONF", "报表": "RPT",
                    }
                    abbr = abbr_map.get(module, "".join(c for c in module if c.isalpha())[:4].upper() or "MISC")

                    db.add(AutomationScript(
                        id=str(uuid.uuid4()),
                        project_id=project_id,
                        test_case_id=(matched_tc or test_cases[0]).id,
                        script_code=f"SC_{abbr}_{total_saved + 1:03d}",
                        script_type=ai_script.get("scriptType", "UI"),
                        framework=ai_script.get("framework", "Playwright"),
                        language=ai_script.get("language", "Python"),
                        code=ai_script.get("code", "# AI generation failed"),
                        status="待执行",
                        generated_by_ai=True,
                    ))
                await db.commit()
                total_saved += len(batch)
                yield f"data: {json.dumps({'event': 'progress', 'count': total_saved, 'batch': len(batch)}, ensure_ascii=False)}\n\n"

            yield f"data: {json.dumps({'event': 'done', 'total': total_saved}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.exception(f"stream_generate_scripts failed: project_id={project_id}")
            yield f"data: {json.dumps({'event': 'error', 'message': _friendly_error(e, '脚本生成')}, ensure_ascii=False)}\n\n"


@router.get("/projects/{project_id}/ai/stream-scripts")
async def stream_scripts(
    project_id: str,
    user: dict = Depends(get_current_user_sse),
    db: AsyncSession = Depends(get_db),
):
    await verify_project_owner(db, project_id, user["sub"])

    # 检查测试用例
    tc_result = await db.execute(
        select(TestCase).where(
            TestCase.project_id == project_id,
            TestCase.automation == "适合"
        )
    )
    if not tc_result.scalars().first():
        raise HTTPException(status_code=400, detail="没有适合自动化的测试用例")

    # 检查配置
    config_check = await check_config_for_task("脚本生成", user["sub"])
    if not config_check["configured"]:
        raise HTTPException(status_code=400, detail=config_check["message"])

    return StreamingResponse(
        _stream_generate_scripts(project_id, user["sub"]),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
