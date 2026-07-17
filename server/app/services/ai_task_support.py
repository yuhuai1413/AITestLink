"""AI task parsing helpers kept independent from HTTP routing."""

import os
from datetime import datetime, timezone

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.ai_task import AITask
from app.models.file_asset import FileAsset


async def update_task_status(
    db: AsyncSession,
    task_id: str,
    status: str,
    error: str | None = None,
) -> None:
    result = await db.execute(select(AITask).where(AITask.id == task_id))
    task = result.scalar_one_or_none()
    if task:
        task.status = status
        task.finished_at = datetime.now(timezone.utc)
        if error:
            task.error_message = error[:2000]
        await db.commit()


def normalize_automation(value: object) -> str:
    """Normalize model automation hints to the persisted Chinese enum."""
    if value is None:
        return "否"
    if isinstance(value, bool):
        return "是" if value else "否"
    return "是" if str(value).strip().lower() in ("true", "1", "适合", "yes", "是") else "否"


def friendly_error(err: Exception, task_type: str = "") -> str:
    """Convert transport/model exceptions to stable user-facing messages."""
    message = str(err)
    if isinstance(err, httpx.ConnectTimeout):
        return "模型服务连接超时，请检查网络连接或模型配置中的 API 地址"
    if isinstance(err, httpx.ReadTimeout):
        return "模型服务响应超时，请稍后重试（模型生成内容较多时可能需要更长时间）"
    if isinstance(err, httpx.WriteTimeout):
        return "模型服务发送数据超时，请稍后重试"
    if isinstance(err, httpx.PoolTimeout):
        return "模型服务连接池超时，请稍后重试"
    if isinstance(err, httpx.TimeoutException):
        return "模型服务响应超时，请检查网络连接后重试"
    if isinstance(err, httpx.ConnectError):
        return "无法连接到模型服务，请检查网络连接或模型配置中的 API 地址"

    response = getattr(err, "response", None)
    code = getattr(response, "status_code", 0) if response is not None else 0
    if code == 400:
        return f"{task_type or '请求'}参数错误（400），请检查输入数据后重试"
    if code in (401, 403):
        return "API Key 无效或无权限，请在模型配置中检查 API Key"
    if code == 404:
        return "模型服务地址不可用（404），请在模型配置中检查 API 地址是否正确"
    if code == 429:
        return "模型服务请求过于频繁（限流），请稍后重试"
    if code >= 500:
        return "模型服务暂时不可用，请稍后重试"

    if "404" in message or "Not Found" in message:
        return "模型服务地址不可用（404），请在模型配置中检查 API 地址是否正确"
    if any(token in message for token in ("401", "Unauthorized", "403", "Forbidden")):
        return "API Key 无效或无权限，请在模型配置中检查 API Key"
    if any(token in message for token in ("429", "Too Many Requests", "Rate limit")):
        return "模型服务请求过于频繁（限流），请稍后重试"
    if any(token in message for token in ("500", "502", "503", "Internal Server", "Bad Gateway", "Service Unavailable")):
        return "模型服务暂时不可用，请稍后重试"
    if any(token in message for token in ("Connect", "connect", "ConnectionRefused")):
        return "无法连接到模型服务，请检查网络连接或模型配置中的 API 地址"
    if "Timeout" in message or "timeout" in message:
        return "模型服务响应超时，请检查网络连接后重试"
    if "SSL" in message or "ssl" in message:
        return "SSL 连接错误，请检查网络环境"
    if "JSONDecodeError" in message or "json" in message.lower():
        return "模型返回的数据格式异常，无法解析，请稍后重试"
    if "结果结构校验失败" in message:
        if task_type == "测试点生成" and "requirementId" in message:
            return (
                "测试点生成失败：模型返回的测试点缺少「关联需求ID」。"
                "系统无法判断测试点属于哪条需求，因此已停止写入，避免数据关联错误。"
                "请重新生成；如果连续失败，请检查「测试点生成」模型配置或提示词。"
            )
        return f"{task_type or 'AI任务'}失败：模型返回结果不符合系统要求。{message[:160]}"
    if "KeyError" in message:
        return "模型返回数据结构异常，请稍后重试"
    if "配置不存在" in message or "模型配置" in message or "请先" in message:
        return message

    label = task_type or "任务"
    return f"{label}失败：{message[:200]}" if message else f"{label}失败：未知错误，请稍后重试"


MODULE_ABBR: dict[str, str] = {
    "用户管理": "USER", "用户": "USER", "登录": "LOGIN", "注销": "LOGOUT", "注册": "SIGNUP",
    "权限": "PERM", "角色": "ROLE", "账号": "ACCOUNT", "密码": "PWD", "退出": "LOGOUT",
    "客户管理": "CUST", "客户": "CUST", "线索": "LEAD", "商机": "OPP", "合同": "CONTRACT",
    "订单": "ORDER", "交易": "TRADE", "支付": "PAY", "退款": "REFUND", "购物车": "CART",
    "商品": "GOODS", "产品": "PROD", "库存": "STOCK", "分类": "CATEGORY", "供应商": "SUPPLIER",
    "财务": "FINANCE", "发票": "INVOICE", "报销": "REIMBURSE", "预算": "BUDGET", "账单": "BILL",
    "审批": "APPROVE", "流程": "FLOW", "工作流": "WORKFLOW", "通知": "NOTIF", "消息": "MSG",
    "报表": "REPORT", "统计": "STAT", "分析": "ANALYSIS", "图表": "CHART", "仪表盘": "DASH",
    "文件": "FILE", "文档": "DOC", "模板": "TPL", "附件": "ATTACH", "上传": "UPLOAD",
    "系统": "SYS", "设置": "SETTING", "配置": "CONFIG", "参数": "PARAM", "日志": "LOG",
    "搜索": "SEARCH", "筛选": "FILT", "排序": "SORT", "表格": "TABLE", "表单": "FORM",
    "列表": "LIST", "详情": "DETAIL", "编辑": "EDIT", "新增": "ADD", "删除": "DEL", "首页": "HOME",
    "测试": "TEST", "缺陷": "BUG", "需求": "REQ", "脚本": "SCRIPT", "用例": "CASE",
    "执行": "EXEC", "评审": "REVIEW", "生成": "GEN", "回归": "REG", "冒烟": "SMOKE",
    "接口": "API", "性能": "PERF", "安全": "SEC", "兼容": "COMPAT", "自动化": "AUTO",
    "核心": "CORE", "基础": "BASE", "公共": "COMMON", "通用": "COMMON", "状态": "STATUS",
    "版本": "VERSION", "发布": "RELEASE", "部署": "DEPLOY", "缓存": "CACHE", "任务": "TASK",
    "数据权限": "DATAPERM",
}


def to_eng_abbr(module: str) -> str:
    if module in MODULE_ABBR:
        return MODULE_ABBR[module]
    for chinese, english in MODULE_ABBR.items():
        if len(chinese) >= 2 and chinese in module:
            remaining = module.replace(chinese, "", 1)
            return f"{english}_{MODULE_ABBR[remaining]}" if remaining in MODULE_ABBR else english
    ascii_chars = [char.upper() for char in module if char.isascii() and char.isalpha()]
    return "".join(ascii_chars[:4]) if ascii_chars else "MOD"


def module_counter(items: list[dict], prefix: str) -> list[tuple[str, str]]:
    counters: dict[str, int] = {}
    result: list[tuple[str, str]] = []
    for item in items:
        module = item["module"]
        abbreviation = to_eng_abbr(module)
        counters[abbreviation] = counters.get(abbreviation, 0) + 1
        result.append((f"{prefix}_{abbreviation}_{counters[abbreviation]:03d}", module))
    return result


def read_file_content(file_obj: FileAsset) -> str:
    if not file_obj.storage_path or not os.path.exists(file_obj.storage_path):
        return ""
    extension = file_obj.name.rsplit(".", 1)[-1].lower() if "." in file_obj.name else ""
    try:
        if extension in ("docx", "doc"):
            from docx import Document
            document = Document(file_obj.storage_path)
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for index, table in enumerate(document.tables):
                rows = [" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()) for row in table.rows]
                if any(rows):
                    parts.append(f"[表格{index + 1}]\n" + "\n".join(filter(None, rows)))
            return "\n".join(parts)
        if extension in ("xlsx", "xls"):
            import openpyxl
            workbook = openpyxl.load_workbook(file_obj.storage_path, read_only=True, data_only=True)
            parts = []
            for sheet_name in workbook.sheetnames:
                rows = [" | ".join(str(cell).strip() for cell in row if cell is not None and str(cell).strip()) for row in workbook[sheet_name].iter_rows(values_only=True)]
                if any(rows):
                    parts.append(f"Sheet: {sheet_name}\n" + "\n".join(filter(None, rows)))
            workbook.close()
            return "\n".join(parts)
        if extension == "pdf":
            import PyPDF2
            return "\n".join(page.extract_text() or "" for page in PyPDF2.PdfReader(file_obj.storage_path).pages)
        if extension in ("md", "txt", "json", "yaml", "yml", "csv"):
            with open(file_obj.storage_path, "r", errors="ignore") as file_handle:
                return file_handle.read()
        return f"(不支持的格式: {extension})"
    except Exception as error:
        return f"(读取失败: {str(error)[:100]})"
