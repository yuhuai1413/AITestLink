from __future__ import annotations

import base64
import os
import tempfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable


def _safe(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _fmt_date(value: Any) -> str:
    if not value:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    text = str(value)
    return text[:10] if len(text) >= 10 else text


def _format_steps(text: str) -> str:
    text = (text or "").replace("\\n", "\n").strip()
    if not text:
        return ""
    if "\n" in text:
        return text
    parts = [p.strip() for p in text.replace("；", ";").split(";") if p.strip()]
    if len(parts) <= 1:
        return text
    return "\n".join(parts)


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    original = paragraph.text
    replaced = original
    for key, value in replacements.items():
        replaced = replaced.replace(key, value)
    if replaced == original:
        return

    # Prefer run-level replacement to keep formatting where placeholders are not
    # split across runs. Fall back to first-run replacement for split cases.
    changed_by_run = False
    for run in paragraph.runs:
        run_text = run.text
        new_text = run_text
        for key, value in replacements.items():
            new_text = new_text.replace(key, value)
        if new_text != run_text:
            run.text = new_text
            changed_by_run = True
    if changed_by_run and paragraph.text == replaced:
        return

    paragraph.runs[0].text = replaced
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_cell(cell, replacements: dict[str, str]) -> None:
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, replacements)


def _set_cell_text(cell, value: Any) -> None:
    text = _safe(value)
    if not cell.paragraphs:
        cell.text = text
        return
    cell.paragraphs[0].text = text
    for paragraph in cell.paragraphs[1:]:
        paragraph.text = ""


def _clear_row(row) -> None:
    for cell in row.cells:
        _set_cell_text(cell, "")


def _clone_row_after(table, row_index: int):
    source = table.rows[row_index]._tr
    new_tr = deepcopy(source)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _header_texts(table) -> list[str]:
    if not table.rows:
        return []
    return [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]


def _all_table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _is_placeholder_row(row) -> bool:
    text = " ".join(cell.text for cell in row.cells)
    return "[" in text or "{{" in text


def _data_start_index(table) -> int:
    if len(table.rows) <= 1:
        return 1
    return 1


def _fill_repeating_rows(table, items: Iterable[Any], value_for_header) -> None:
    headers = _header_texts(table)
    data_start = _data_start_index(table)
    items = list(items)
    if not items:
        for row in table.rows[data_start:]:
            _clear_row(row)
        return

    while len(table.rows) < data_start + len(items):
        _clone_row_after(table, data_start if data_start < len(table.rows) else len(table.rows) - 1)

    for offset, item in enumerate(items):
        row = table.rows[data_start + offset]
        for col_index, cell in enumerate(row.cells):
            header = headers[col_index] if col_index < len(headers) else ""
            _set_cell_text(cell, value_for_header(item, header, col_index))

    for row in table.rows[data_start + len(items):]:
        if _is_placeholder_row(row) or not any(cell.text.strip() for cell in row.cells):
            _clear_row(row)


def _requirement_value(item, header: str, index: int) -> str:
    if "需求标识" in header or "需求编号" in header:
        return _safe(getattr(item, "req_id", "")) or _safe(getattr(item, "id", ""))
    if "模块" in header:
        return _safe(getattr(item, "module", ""))
    if "测试项" in header or "功能" in header or "需求" in header:
        return _safe(getattr(item, "feature", "")) or _safe(getattr(item, "rule", ""))
    if "来源" in header:
        return _safe(getattr(item, "source", ""))
    if "风险" in header or "优先级" in header:
        return _safe(getattr(item, "risk", ""))
    if "说明" in header or "规则" in header or "描述" in header:
        return _safe(getattr(item, "rule", ""))
    values = [
        getattr(item, "req_id", ""),
        getattr(item, "feature", ""),
        getattr(item, "rule", ""),
    ]
    return _safe(values[index]) if index < len(values) else ""


def _point_value(item, header: str, index: int) -> str:
    if "需求标识" in header:
        return _safe(getattr(item, "requirement_id", ""))
    if "测试点编号" in header or "测试点标识" in header:
        return _safe(getattr(item, "point_code", "")) or _safe(getattr(item, "id", ""))
    if "测试项" in header or "测试点" in header or "标题" in header:
        return _safe(getattr(item, "title", ""))
    if "测试类型" in header or header == "类型":
        return _safe(getattr(item, "type", ""))
    if "模块" in header:
        return _safe(getattr(item, "module", ""))
    if "优先级" in header:
        return _safe(getattr(item, "priority", ""))
    if "描述" in header or "说明" in header:
        return _safe(getattr(item, "description", ""))
    values = [
        getattr(item, "point_code", ""),
        getattr(item, "title", ""),
        getattr(item, "type", ""),
    ]
    return _safe(values[index]) if index < len(values) else ""


def _case_value(item, header: str, index: int) -> str:
    if "测试用例标识" in header or "用例标识" in header or "用例编号" in header:
        return _safe(getattr(item, "case_code", ""))
    if "用例标题" in header or "标题" in header:
        return _safe(getattr(item, "title", ""))
    if "测试点" in header:
        return _safe(getattr(item, "feature", "")) or _safe(getattr(item, "title", ""))
    if "模块" in header:
        return _safe(getattr(item, "module", ""))
    if "优先级" in header:
        return _safe(getattr(item, "priority", ""))
    if "前置" in header or "预置" in header:
        return _safe(getattr(item, "precondition", ""))
    if "步骤" in header or "测试过程" in header:
        return _format_steps(_safe(getattr(item, "steps", "")))
    if "测试数据" in header:
        return _safe(getattr(item, "test_data", ""))
    if "预期" in header:
        return _safe(getattr(item, "expected_result", ""))
    if "实测" in header or "实际" in header:
        return _safe(getattr(item, "actual_result", ""))
    if "结果判定" in header or header == "结果":
        return _safe(getattr(item, "passed", "")) or "未执行"
    if "备注" in header:
        return _safe(getattr(item, "remark", ""))
    values = [
        getattr(item, "case_code", ""),
        getattr(item, "expected_result", ""),
        getattr(item, "actual_result", ""),
        getattr(item, "passed", ""),
        getattr(item, "remark", ""),
    ]
    return _safe(values[index]) if index < len(values) else ""


def _fill_case_detail_table(table, case, project_name: str) -> None:
    replacements = {
        "[项目名称]": project_name,
        "[待填写]": "",
        "[测试用例标识]": _safe(getattr(case, "case_code", "")),
        "[测试用例名称]": _safe(getattr(case, "title", "")),
        "[测试项]": _safe(getattr(case, "feature", "")) or _safe(getattr(case, "title", "")),
        "[测试说明]": _safe(getattr(case, "precondition", "")),
        "[测试过程]": _format_steps(_safe(getattr(case, "steps", ""))),
        "[预期结果]": _safe(getattr(case, "expected_result", "")),
        "[测试数据]": _safe(getattr(case, "test_data", "")),
    }
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, replacements)

    # Many uploaded test-spec templates contain empty reserved rows instead of
    # named placeholders. Fill the common 11-row layout deterministically.
    labels = [
        ("项目名称", project_name, "用例标识", _safe(getattr(case, "case_code", ""))),
        ("用例标题", _safe(getattr(case, "title", "")), "优先级", _safe(getattr(case, "priority", ""))),
        ("所属模块", _safe(getattr(case, "module", "")), "测试端", _safe(getattr(case, "target_platform", ""))),
        ("前置条件", _safe(getattr(case, "precondition", "")), "角色", _safe(getattr(case, "required_role", ""))),
        ("测试地址", _safe(getattr(case, "test_url", "")), "测试数据", _safe(getattr(case, "test_data", ""))),
        ("测试步骤", _format_steps(_safe(getattr(case, "steps", ""))), "预期结果", _safe(getattr(case, "expected_result", ""))),
        ("实际结果", _safe(getattr(case, "actual_result", "")), "结果判定", _safe(getattr(case, "passed", "")) or "未执行"),
    ]
    for row_index, values in enumerate(labels):
        if row_index >= len(table.rows):
            break
        row = table.rows[row_index]
        for cell_index, value in enumerate(values):
            if cell_index < len(row.cells):
                _set_cell_text(row.cells[cell_index], value)


def _classify_table(table) -> str:
    headers = _header_texts(table)
    joined = " ".join(headers) + " " + _all_table_text(table)
    if "测试用例标识" in joined or "用例标识" in joined:
        if len(table.rows) > 3 and len(table.columns) >= 4 and "项目名称" in joined:
            return "case_detail"
        return "test_cases"
    if "需求标识" in joined and "测试项" in joined:
        return "test_points" if "测试类型" in joined else "requirements"
    if "发布日期" in joined and "更改描述" in joined:
        return "change_log"
    if "环境类型" in joined and ("当前配置" in joined or "计划配置" in joined):
        return "environment"
    return "unknown"


def render_docx_from_template(
    template_path: str | Path,
    *,
    template_name: str,
    project,
    requirements: list[Any],
    test_points: list[Any],
    test_cases: list[Any],
) -> dict[str, Any]:
    from docx import Document

    path = Path(template_path)
    doc = Document(str(path))
    project_name = _safe(getattr(project, "name", "")) or "未命名项目"
    today = datetime.now().strftime("%Y-%m-%d")
    requirement_code_by_id = {
        _safe(getattr(item, "id", "")): (_safe(getattr(item, "req_id", "")) or _safe(getattr(item, "id", "")))
        for item in requirements
    }
    test_point_code_by_id = {
        _safe(getattr(item, "id", "")): (_safe(getattr(item, "point_code", "")) or _safe(getattr(item, "id", "")))
        for item in test_points
    }

    def point_value(item, header: str, index: int) -> str:
        if "需求标识" in header:
            raw = _safe(getattr(item, "requirement_id", ""))
            return requirement_code_by_id.get(raw, raw)
        return _point_value(item, header, index)

    def case_value(item, header: str, index: int) -> str:
        if "需求标识" in header:
            raw = _safe(getattr(item, "requirement_id", ""))
            return requirement_code_by_id.get(raw, raw)
        if "测试点标识" in header or "测试点编号" in header:
            raw = _safe(getattr(item, "test_point_id", ""))
            return test_point_code_by_id.get(raw, raw)
        return _case_value(item, header, index)

    replacements = {
        "[项目名称]": project_name,
        "[软件名称]": project_name,
        "[软件标识]": _safe(getattr(project, "id", "")),
        "[软件版本]": "V1.0",
        "[文件编号]": f"DOC-{datetime.now().strftime('%Y%m%d')}",
        "[发布日期]": today,
        "[页数]": "",
        "[更改描述]": "根据项目需求、测试点和测试用例自动生成",
        "[填写软件用途、业务目标和适用范围。]": _safe(getattr(project, "description", "")) or "根据项目需求文档覆盖核心业务流程、异常场景和数据校验。",
        "{{project.name}}": project_name,
        "{{project.testType}}": _safe(getattr(project, "test_type", "")),
        "{{project.description}}": _safe(getattr(project, "description", "")),
        "{{summary.requirementCount}}": str(len(requirements)),
        "{{summary.testPointCount}}": str(len(test_points)),
        "{{summary.testCaseCount}}": str(len(test_cases)),
        "{{today}}": today,
    }

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    case_detail_index = 0
    for table in doc.tables:
        kind = _classify_table(table)
        if kind == "requirements":
            _fill_repeating_rows(table, requirements, _requirement_value)
        elif kind == "test_points":
            source = test_points or requirements
            value_fn = point_value if test_points else _requirement_value
            _fill_repeating_rows(table, source, value_fn)
        elif kind == "test_cases":
            _fill_repeating_rows(table, test_cases, case_value)
        elif kind == "case_detail":
            if case_detail_index < len(test_cases):
                _fill_case_detail_table(table, test_cases[case_detail_index], project_name)
            else:
                for row in table.rows:
                    _clear_row(row)
            case_detail_index += 1
        elif kind == "change_log":
            _fill_repeating_rows(
                table,
                [{"date": today, "desc": "系统根据项目数据生成文档"}],
                lambda item, header, index: item["date"] if "日期" in header else item["desc"],
            )
        elif kind == "environment":
            _fill_repeating_rows(
                table,
                [{"type": "测试环境", "current": "以项目环境配置为准", "plan": "按测试计划执行"}],
                lambda item, header, index: item["type"] if "环境" in header else (item["current"] if "当前" in header else item["plan"]),
            )

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    try:
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as f:
            docx_bytes = f.read()
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass

    title = template_name or path.stem
    return {
        "documentType": title,
        "title": f"{project_name}-{title}",
        "content": f"已基于模板「{path.name}」生成 Word 文档。需求 {len(requirements)} 条，测试点 {len(test_points)} 个，测试用例 {len(test_cases)} 条。",
        "docxBase64": base64.b64encode(docx_bytes).decode(),
        "docxFileName": f"{project_name}-{title}.docx",
        "metadata": {
            "templateFile": path.name,
            "generationMode": "template-render",
            "requirementCount": len(requirements),
            "testPointCount": len(test_points),
            "testCaseCount": len(test_cases),
        },
    }
