from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.services.export_format import format_api_datetime


PLACEHOLDER_RE = re.compile(r"\[[^\[\]\n]{1,80}\]|\{\{[^{}\n]{1,120}\}\}")


def sha256_file(path: str | Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_placeholders(text: str) -> list[str]:
    seen: set[str] = set()
    placeholders: list[str] = []
    for match in PLACEHOLDER_RE.findall(text or ""):
        if match not in seen:
            seen.add(match)
            placeholders.append(match)
    return placeholders


def _table_kind(headers: list[str], all_text: str) -> str:
    joined = " ".join(headers) + " " + all_text
    if "测试用例标识" in joined or "用例标识" in joined:
        return "test_cases"
    if "需求标识" in joined and "测试项" in joined:
        return "test_points" if "测试类型" in joined else "requirements"
    if "严重程度" in joined and "数量" in joined and "百分比" in joined:
        return "defect_summary"
    if "环境类型" in joined and ("当前配置" in joined or "计划配置" in joined):
        return "environment"
    if "发布日期" in joined and "更改描述" in joined:
        return "change_log"
    if "缩写" in joined and "英文全称" in joined:
        return "glossary"
    return "unknown"


def parse_docx_template(path: str | Path) -> dict[str, Any]:
    """Extract a compact, cached structure from a Word template.

    The structure is intentionally metadata-only. Generation still starts from
    the original docx file so styles, margins, headers, footers and table
    formatting remain owned by the uploaded template.
    """
    from docx import Document

    file_path = Path(path)
    file_hash = sha256_file(file_path)
    doc = Document(str(file_path))

    paragraphs: list[dict[str, Any]] = []
    placeholders: list[str] = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text or ""
        ph = extract_placeholders(text)
        if text.strip() or ph:
            paragraphs.append({
                "index": index,
                "text": text.strip(),
                "style": paragraph.style.name if paragraph.style else "",
                "placeholders": ph,
            })
            placeholders.extend(ph)

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables):
        headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells] if table.rows else []
        sample_rows: list[list[str]] = []
        table_placeholders: list[str] = []
        for row in table.rows[:4]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            sample_rows.append(values)
            for value in values:
                table_placeholders.extend(extract_placeholders(value))
        all_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        kind = _table_kind(headers, all_text)
        tables.append({
            "index": table_index,
            "kind": kind,
            "rowCount": len(table.rows),
            "columnCount": len(table.columns),
            "headers": headers,
            "sampleRows": sample_rows,
            "placeholders": list(dict.fromkeys(table_placeholders)),
        })
        placeholders.extend(table_placeholders)

    unique_placeholders = list(dict.fromkeys(placeholders))
    table_kinds: dict[str, int] = {}
    for table in tables:
        table_kinds[table["kind"]] = table_kinds.get(table["kind"], 0) + 1

    return {
        "version": 1,
        "fileName": file_path.name,
        "fileHash": file_hash,
        "parsedAt": format_api_datetime(datetime.now(timezone.utc)),
        "paragraphCount": len(doc.paragraphs),
        "tableCount": len(doc.tables),
        "placeholders": unique_placeholders,
        "tableKinds": table_kinds,
        "paragraphs": paragraphs[:80],
        "tables": tables,
        "capabilities": {
            "replaceProjectFields": True,
            "fillRequirementTables": any(t["kind"] == "requirements" for t in tables),
            "fillTestPointTables": any(t["kind"] == "test_points" for t in tables),
            "fillTestCaseTables": any(t["kind"] == "test_cases" for t in tables),
        },
    }


def dumps_structure(structure: dict[str, Any]) -> str:
    return json.dumps(structure, ensure_ascii=False, separators=(",", ":"))
