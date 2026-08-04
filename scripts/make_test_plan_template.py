"""把测试计划 docx 处理成带 {{xxx}} 占位符的模板文件。

只做文本替换（run 级别），不动任何格式（字体/字号/段落/表格样式全部保持原样）。
表 12/13（需求追踪表）的示例数据行清空为占位符行，让渲染器填充真实数据。
"""
from docx import Document

SRC = "docs/测试文档模板/客户关系管理系统软件测试计划.docx"
DST = "docs/测试文档模板/软件测试计划-模板.docx"

# 项目相关内容 → 占位符映射（只替换项目特定的值，通用文本不动）
REPLACEMENTS = {
    # 软件名称
    "客户关系管理系统": "{{software.name}}",
    # 项目名称
    "营销数字化提升项目": "{{project.name}}",
    # 软件标识
    "CRM": "{{software.code}}",
    # 软件版本
    "V1.0": "{{software.version}}",
    # 开发方
    "宁夏世纪信通信息安全有限公司": "{{dev.company}}",
    # 需方/用户
    "宁夏伊品生物科技股份有限公司": "{{client.company}}",
    # 文件编号
    "SJXT-YF-CJ-001": "{{doc.number}}",
    # 开发启动时间
    "2025年9月": "{{project.startDate}}",
    # 维护期限
    "18个月": "{{maintenance.months}}",
}


def replace_in_runs(paragraph):
    """在 run 级别做文本替换，保持格式。处理占位符跨 run 的情况。"""
    # 先尝试整段文本替换（占位符可能跨多个 run）
    original = paragraph.text
    replaced = original
    for old, new in REPLACEMENTS.items():
        replaced = replaced.replace(old, new)
    if replaced == original:
        return False
    # 尝试逐 run 替换
    changed = False
    for run in paragraph.runs:
        rt = run.text
        new_rt = rt
        for old, new in REPLACEMENTS.items():
            new_rt = new_rt.replace(old, new)
        if new_rt != rt:
            run.text = new_rt
            changed = True
    # 如果逐 run 没替换成功（占位符跨 run），把整段结果放到第一个 run
    if not changed and paragraph.runs:
        for run in paragraph.runs[1:]:
            run.text = ""
        paragraph.runs[0].text = replaced
        changed = True
    return changed


def clear_table_data_rows(table, header_count=1):
    """清空表格的数据行（保留表头），用于需求追踪表等示例数据。"""
    rows = list(table.rows)
    if len(rows) <= header_count:
        return
    # 第 header_count 行作为占位符模板行（清空内容）
    template_row = rows[header_count]
    for cell in template_row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
    # 删除多余的数据行（只保留表头 + 1 个占位符行）
    for row in rows[header_count + 1:]:
        row._element.getparent().remove(row._element)


def main():
    doc = Document(SRC)

    # 1. 段落文本替换
    para_count = 0
    for p in doc.paragraphs:
        if replace_in_runs(p):
            para_count += 1

    # 2. 表格文本替换
    table_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_runs(p):
                        table_count += 1

    # 3. 表 12/13（需求追踪表，130 行 x 2 列）清空示例数据
    #    表 12: 测试项 → 需求标识；表 13: 需求标识 → 测试项
    cleared_tables = 0
    for table in doc.tables:
        if len(table.rows) >= 100 and len(table.columns) == 2:
            # 大概率是需求追踪表（行数多、2 列）
            header = [c.text.strip() for c in table.rows[0].cells]
            if any("测试项" in h or "需求标识" in h or "测试用例" in h for h in header):
                clear_table_data_rows(table, header_count=1)
                cleared_tables += 1

    doc.save(DST)
    print(f"模板已生成: {DST}")
    print(f"  段落替换: {para_count} 处")
    print(f"  表格替换: {table_count} 处")
    print(f"  需求追踪表清空: {cleared_tables} 个")


if __name__ == "__main__":
    main()
