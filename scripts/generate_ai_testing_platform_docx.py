from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AI软件测试平台产品方案.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "C9D2DC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_inch: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths_inch):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(round(width * 1440)))


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=4, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=24, color=INK, bold=True)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=13, color=MUTED)


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_heading("", level=level)
    if level == 1:
        set_paragraph_spacing(p, before=16, after=8, line=1.10)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(p, before=12, after=6, line=1.10)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, before=8, after=4, line=1.10)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.10)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.167)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4, line=1.167)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table, color="D9E1EA")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.10)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.10)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_FILL)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.0)
        r = p.add_run(header)
        set_run_font(r, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_shading(cells[i], WHITE)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=0, line=1.05)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_flow(doc: Document) -> None:
    steps = [
        "创建测试项目",
        "导入需求/原型/接口文档",
        "AI 解析需求",
        "生成待确认问题",
        "生成测试点",
        "生成测试用例",
        "人工评审与调整",
        "准备测试数据",
        "执行测试",
        "生成缺陷",
        "回归验证",
        "输出测试报告",
        "沉淀测试资产",
    ]
    add_callout(doc, "标准测试闭环", " → ".join(steps))


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
    header = section.header.paragraphs[0]
    header.text = "AI 软件测试平台产品方案"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0, line=1.0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "内部方案草案"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)
    return doc


def build() -> None:
    doc = setup_doc()

    add_title(doc, "AI 软件测试平台产品方案")
    add_subtitle(doc, "基于标准软件测试流程与 AI 能力的产品、MVP、技术架构和数据模型设计")

    meta_rows = [
        ["文档类型", "产品方案 / 技术方案"],
        ["适用对象", "测试经理、测试工程师、自动化测试工程师、研发负责人、产品负责人"],
        ["版本", "V1.0"],
        ["日期", "2026-07-01"],
    ]
    add_table(doc, ["项目", "内容"], meta_rows, [1.5, 5.0])

    add_callout(
        doc,
        "核心定位",
        "平台不是单一的 AI 生成用例工具，而是围绕需求解析、测试设计、测试执行、缺陷分析、自动化测试、报告生成和资产沉淀构建的测试全流程工作台。",
    )

    add_heading(doc, "1. 建设目标")
    add_bullets(
        doc,
        [
            "降低需求分析和测试设计的重复劳动。",
            "提升测试点、测试用例、异常场景和边界场景覆盖率。",
            "建立需求、测试点、测试用例、执行结果、缺陷和报告之间的可追溯关系。",
            "通过 AI 辅助生成自动化脚本，提高回归测试效率。",
            "通过 AI 分析执行日志、截图和接口响应，提升失败定位效率。",
            "沉淀历史需求、用例、缺陷、脚本和经验，形成企业级测试知识库。",
        ],
    )

    add_heading(doc, "2. 用户角色")
    add_table(
        doc,
        ["角色", "主要职责", "平台诉求"],
        [
            ["测试经理", "创建项目、制定测试计划、把控质量风险、输出报告", "需要项目质量全貌、测试进度、缺陷趋势和上线建议"],
            ["测试工程师", "分析需求、设计用例、执行测试、提交缺陷", "需要高质量测试点、可执行用例、便捷执行记录和缺陷生成"],
            ["自动化测试工程师", "维护自动化脚本、执行回归、分析失败原因", "需要脚本生成、执行编排、日志分析和定位器维护能力"],
            ["产品/业务人员", "提供需求、确认问题、评审测试覆盖", "需要查看需求覆盖、待确认问题和测试范围"],
            ["开发人员", "修复缺陷、分析失败、参与回归验证", "需要准确复现步骤、日志、截图和影响范围"],
        ],
        [1.25, 2.55, 2.70],
    )

    add_heading(doc, "3. 标准业务流程")
    add_flow(doc)
    add_para(doc, "平台设计中要保留人工评审节点。AI 生成的需求分析、测试点、测试用例和缺陷建议不能直接作为最终结论，必须经过测试人员确认。")

    add_heading(doc, "4. 测试类型支持")
    add_table(
        doc,
        ["测试类型", "典型场景", "平台处理方式"],
        [
            ["首轮全量测试", "新系统首次测试、无历史用例", "从需求解析开始，完整生成测试点、用例、数据、报告"],
            ["全量复用测试", "有旧用例但需要重新建立覆盖", "对比旧用例和新需求，补充遗漏、删除过期项"],
            ["回归测试", "版本变更、缺陷修复后验证", "根据变更影响筛选已有用例，补充必要新用例"],
            ["增量测试", "新增或修改局部功能", "识别差异、评估影响范围、生成增量用例"],
            ["探索性测试", "需求不完整或需要经验驱动发现问题", "生成探索大纲，记录发现，沉淀为正式用例"],
        ],
        [1.4, 2.3, 2.8],
    )

    add_heading(doc, "5. 核心功能模块")
    add_table(
        doc,
        ["模块", "作用", "AI 能力"],
        [
            ["首页驾驶舱", "展示质量状态、执行进度、缺陷趋势和风险提醒", "自动汇总质量风险和待处理事项"],
            ["项目空间", "管理项目资料、测试计划、用例、缺陷和报告", "按项目上下文组织 AI 任务和资产"],
            ["需求解析", "上传并解析需求、原型、接口文档和变更说明", "提取模块、功能点、业务规则、风险点和待确认问题"],
            ["测试设计", "生成测试点、测试用例和测试数据", "覆盖正常、异常、边界、权限、数据一致性和状态流转"],
            ["用例管理", "维护、评审、复用和导出测试用例", "识别重复用例、遗漏场景和过期用例"],
            ["执行中心", "支持手工执行和自动化执行记录", "分析执行失败原因并推荐处理方式"],
            ["自动化中心", "生成、维护和执行 UI/API 自动化脚本", "生成脚本、修复定位器、分析失败日志"],
            ["缺陷中心", "创建、流转、关联和分析缺陷", "生成标准缺陷单和严重程度建议"],
            ["报告中心", "生成日报、回归报告和版本测试报告", "总结测试范围、缺陷分布、遗留风险和上线建议"],
            ["知识库", "沉淀历史需求、用例、缺陷、脚本和经验", "相似检索、用例复用和回归推荐"],
        ],
        [1.25, 2.75, 2.50],
    )

    add_heading(doc, "6. AI Agent 设计")
    add_table(
        doc,
        ["Agent", "职责"],
        [
            ["需求解析 Agent", "解析文档、提取功能点、识别业务规则和待确认问题"],
            ["测试设计 Agent", "生成测试点、测试用例和测试数据"],
            ["用例评审 Agent", "检查遗漏、重复、不可执行和预期不清"],
            ["自动化 Agent", "生成和维护自动化脚本"],
            ["执行分析 Agent", "分析失败日志、截图、接口响应和 trace"],
            ["缺陷生成 Agent", "生成标准缺陷单"],
            ["报告 Agent", "生成日报、回归报告和版本测试报告"],
            ["回归推荐 Agent", "根据变更影响推荐需要执行的用例集"],
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "7. MVP 范围")
    add_para(doc, "第一版建议聚焦“需求文档 → AI 测试设计 → 用例管理 → Excel 导出”。这是测试人员最容易接受、也最容易产生直接价值的入口。")
    add_heading(doc, "7.1 必须包含", level=2)
    add_bullets(
        doc,
        [
            "项目管理。",
            "文档上传。",
            "AI 需求解析。",
            "AI 测试点生成。",
            "AI 测试用例生成。",
            "用例在线编辑。",
            "用例评审状态。",
            "Excel 导出。",
        ],
    )
    add_heading(doc, "7.2 暂不纳入第一版", level=2)
    add_bullets(
        doc,
        [
            "自动化脚本执行。",
            "缺陷系统深度集成。",
            "CI/CD 编排。",
            "复杂权限模型。",
            "脚本自修复。",
            "质量预测模型。",
        ],
    )

    add_heading(doc, "8. 技术架构建议")
    add_table(
        doc,
        ["层级", "推荐方案", "说明"],
        [
            ["前端", "React + TypeScript + Ant Design", "适合后台管理和复杂表格"],
            ["后端", "FastAPI 或 Spring Boot", "FastAPI 更适合快速集成 AI，Spring Boot 更适合企业级标准化"],
            ["数据库", "PostgreSQL", "支持复杂关系和 JSONB 扩展"],
            ["缓存", "Redis", "用于任务状态、队列状态和临时结果"],
            ["文件存储", "MinIO 或 S3", "存储文档、截图、日志和报告"],
            ["向量库", "pgvector、Milvus 或 Qdrant", "存储需求、用例、缺陷知识向量"],
            ["自动化", "Playwright + pytest", "Web UI 自动化优先"],
            ["接口测试", "pytest + requests/HTTPX", "支持数据驱动和断言"],
            ["任务队列", "Celery、RQ 或 BullMQ", "执行 AI 解析、报告生成和自动化任务"],
        ],
        [1.25, 2.25, 3.0],
    )

    add_heading(doc, "9. 核心数据模型")
    add_para(doc, "平台必须从第一版开始打通“需求 → 测试点 → 测试用例 → 执行记录 → 缺陷 → 回归记录 → 测试报告”的追溯关系。")
    add_table(
        doc,
        ["数据表", "职责"],
        [
            ["project", "项目、版本、测试类型和负责人"],
            ["file_asset", "需求文档、原型、接口文档、截图、日志和报告文件"],
            ["requirement", "AI 解析后的需求模块、功能点、业务规则和待确认问题"],
            ["test_point", "测试点、测试类型、优先级和评审状态"],
            ["test_case", "测试用例主体数据和需求追溯关系"],
            ["test_execution", "手工或自动化执行结果、实际结果和证据"],
            ["bug", "缺陷标题、复现步骤、实际结果、预期结果和状态"],
            ["automation_script", "自动化脚本、框架、仓库路径和关联用例"],
            ["automation_run", "自动化执行记录、通过数、失败数、报告和日志"],
            ["ai_task", "AI 任务类型、输入输出、状态、模型和 Prompt 版本"],
            ["knowledge_item", "企业测试知识库条目和向量检索引用"],
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "10. 演进路线")
    add_table(
        doc,
        ["阶段", "目标", "关键能力"],
        [
            ["MVP 1", "建立 AI 测试设计入口", "文档上传、需求解析、测试点生成、用例生成、Excel 导出"],
            ["MVP 2", "建立测试执行闭环", "用例执行、缺陷生成、执行统计、测试报告"],
            ["MVP 3", "建立自动化能力", "Playwright 脚本生成、API 自动化、自动化任务执行"],
            ["MVP 4", "建立智能质量分析", "失败归因、日志分析、回归推荐、质量风险预测"],
            ["MVP 5", "建立企业知识库", "历史资产复用、相似缺陷检索、组织级测试知识沉淀"],
        ],
        [1.2, 2.2, 3.1],
    )

    add_heading(doc, "11. 结论")
    add_callout(
        doc,
        "产品判断",
        "这个平台的核心竞争力不是简单生成几条测试用例，而是把测试经理、测试工程师和自动化工程师日常重复的分析、设计、执行和归档工作，用 AI 串成一个可追溯、可复用、可持续演进的质量闭环。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()

